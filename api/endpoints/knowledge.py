"""
API para gerenciamento da Base de Conhecimento.
Permite upload, indexação e busca de documentos para a IA.
"""
import os
import uuid
import json
import aiofiles
from datetime import datetime
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel

from database.database import get_db
from database.models import KnowledgeDocument, DocumentType, User
from api.endpoints.auth import get_current_user
from services.vector_store import get_vector_store
from services.document_processor import get_document_processor

router = APIRouter(prefix="/api/knowledge", tags=["Knowledge Base"])

UPLOAD_DIR = "uploads/knowledge"
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_EXTENSIONS = {
    ".pdf": DocumentType.PDF.value,
    ".docx": DocumentType.DOCX.value,
    ".doc": DocumentType.DOCX.value,
    ".txt": DocumentType.TXT.value,
    ".png": DocumentType.IMAGE.value,
    ".jpg": DocumentType.IMAGE.value,
    ".jpeg": DocumentType.IMAGE.value,
}

CATEGORIES = [
    "Estratégias",
    "Produtos",
    "Processos",
    "Compliance",
    "Treinamento",
    "FAQ",
    "Outros"
]


class DocumentResponse(BaseModel):
    id: int
    title: str
    description: Optional[str]
    filename: str
    file_type: str
    file_size: int
    category: Optional[str]
    chunks_count: int
    is_indexed: bool
    index_error: Optional[str]
    created_at: datetime
    
    class Config:
        from_attributes = True


class DocumentListResponse(BaseModel):
    documents: List[DocumentResponse]
    total: int
    categories: List[str]


def extract_text_from_pdf(file_path: str) -> str:
    """Extrai texto de um arquivo PDF."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        raise Exception(f"Erro ao extrair texto do PDF: {str(e)}")


def extract_text_from_docx(file_path: str) -> str:
    """Extrai texto de um arquivo DOCX."""
    try:
        from docx import Document
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise Exception(f"Erro ao extrair texto do DOCX: {str(e)}")


def extract_text_from_txt(file_path: str) -> str:
    """Extrai texto de um arquivo TXT."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read().strip()
    except Exception as e:
        raise Exception(f"Erro ao ler arquivo TXT: {str(e)}")


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    """
    Divide texto em chunks com sobreposição.
    
    Args:
        text: Texto a ser dividido
        chunk_size: Tamanho máximo de cada chunk em caracteres
        overlap: Sobreposição entre chunks
    
    Returns:
        Lista de chunks de texto
    """
    if not text:
        return []
    
    chunks = []
    start = 0
    
    while start < len(text):
        end = start + chunk_size
        
        if end < len(text):
            last_period = text.rfind(".", start, end)
            last_newline = text.rfind("\n", start, end)
            break_point = max(last_period, last_newline)
            
            if break_point > start:
                end = break_point + 1
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap if end < len(text) else end
    
    return chunks


async def index_document_background(doc_id: int, file_path: str, file_type: str, title: str, category: str):
    """
    Processa e indexa um documento em background.
    Extrai texto, divide em chunks e gera embeddings.
    """
    from database.database import SessionLocal
    
    db = SessionLocal()
    try:
        doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if not doc:
            return
        
        try:
            if file_type == DocumentType.PDF.value:
                text = extract_text_from_pdf(file_path)
            elif file_type == DocumentType.DOCX.value:
                text = extract_text_from_docx(file_path)
            elif file_type == DocumentType.TXT.value:
                text = extract_text_from_txt(file_path)
            else:
                doc.is_indexed = False
                doc.index_error = "Tipo de arquivo não suportado para indexação"
                db.commit()
                return
            
            if not text:
                doc.is_indexed = False
                doc.index_error = "Não foi possível extrair texto do documento"
                db.commit()
                return
            
            chunks = chunk_text(text)
            
            if not chunks:
                doc.is_indexed = False
                doc.index_error = "Documento sem conteúdo para indexar"
                db.commit()
                return
            
            vector_store = get_vector_store()
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"doc_{doc_id}_chunk_{i}"
                metadata = {
                    "document_id": doc_id,
                    "document_title": title,
                    "category": category or "Outros",
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                }
                vector_store.add_document(chunk_id, chunk, metadata)
            
            doc.chunks_count = len(chunks)
            doc.is_indexed = True
            doc.index_error = None
            db.commit()
            
            print(f"[KNOWLEDGE] Documento {doc_id} indexado com sucesso: {len(chunks)} chunks")
            
        except Exception as e:
            doc.is_indexed = False
            doc.index_error = str(e)
            db.commit()
            print(f"[KNOWLEDGE] Erro ao indexar documento {doc_id}: {str(e)}")
            
    finally:
        db.close()


async def index_document_smart(doc_id: int, file_path: str, file_type: str, title: str, category: str):
    """
    Processa e indexa um documento usando GPT-4 Vision para análise inteligente.
    Ideal para documentos com tabelas, infográficos e conteúdo visual.
    """
    from database.database import SessionLocal
    
    db = SessionLocal()
    try:
        doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
        if not doc:
            return
        
        try:
            processor = get_document_processor()
            vector_store = get_vector_store()
            
            if file_type in [DocumentType.PDF.value]:
                print(f"[KNOWLEDGE] Processando PDF com IA: {title}")
                processed_data = processor.process_pdf(pdf_path=file_path, document_title=title)
            elif file_type in [DocumentType.IMAGE.value]:
                print(f"[KNOWLEDGE] Processando imagem com IA: {title}")
                processed_data = processor.process_image(image_path=file_path, document_title=title)
            else:
                print(f"[KNOWLEDGE] Tipo {file_type} não suportado para processamento inteligente, usando método padrão")
                await index_document_background(doc_id, file_path, file_type, title, category)
                return
            
            if "error" in processed_data and not processed_data.get("all_facts"):
                doc.is_indexed = False
                doc.index_error = processed_data.get("error", "Erro desconhecido no processamento")
                db.commit()
                return
            
            chunks = processor.generate_indexable_chunks(processed_data)
            
            if not chunks:
                doc.is_indexed = False
                doc.index_error = "Nenhum conteúdo extraído do documento"
                db.commit()
                return
            
            for i, chunk in enumerate(chunks):
                chunk_id = f"doc_{doc_id}_smart_{i}"
                metadata = {
                    "document_id": doc_id,
                    "document_title": title,
                    "category": category or "Outros",
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    "processing_type": "smart",
                    **chunk.get("metadata", {})
                }
                vector_store.add_document(chunk_id, chunk["content"], metadata)
            
            doc.chunks_count = len(chunks)
            doc.is_indexed = True
            doc.index_error = None
            db.commit()
            
            print(f"[KNOWLEDGE] Documento {doc_id} indexado com IA: {len(chunks)} fatos extraídos de {processed_data.get('total_pages', 1)} página(s)")
            
        except Exception as e:
            doc.is_indexed = False
            doc.index_error = f"Erro no processamento inteligente: {str(e)}"
            db.commit()
            print(f"[KNOWLEDGE] Erro ao indexar documento {doc_id} com IA: {str(e)}")
            import traceback
            traceback.print_exc()
            
    finally:
        db.close()


@router.get("/", response_model=DocumentListResponse)
async def list_documents(
    category: Optional[str] = None,
    search: Optional[str] = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Lista todos os documentos da base de conhecimento."""
    if current_user.role not in ["admin", "gestao_rv"]:
        raise HTTPException(status_code=403, detail="Acesso negado")
    
    query = db.query(KnowledgeDocument)
    
    if category:
        query = query.filter(KnowledgeDocument.category == category)
    
    if search:
        query = query.filter(
            KnowledgeDocument.title.ilike(f"%{search}%") |
            KnowledgeDocument.description.ilike(f"%{search}%")
        )
    
    docs = query.order_by(KnowledgeDocument.created_at.desc()).all()
    
    return DocumentListResponse(
        documents=[DocumentResponse.model_validate(doc) for doc in docs],
        total=len(docs),
        categories=CATEGORIES
    )


@router.post("/upload")
async def upload_document(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    title: str = Form(...),
    description: str = Form(None),
    category: str = Form(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Faz upload de um documento e inicia indexação em background.
    Suporta PDF, DOCX, TXT e imagens.
    
    PDFs e imagens utilizam automaticamente processamento inteligente com GPT-4 Vision
    para melhor extração de dados estruturados (tabelas, infográficos).
    """
    if current_user.role not in ["admin", "gestao_rv"]:
        raise HTTPException(status_code=403, detail="Acesso negado")
    
    _, ext = os.path.splitext(file.filename.lower())
    
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Tipo de arquivo não suportado. Permitidos: {', '.join(ALLOWED_EXTENSIONS.keys())}"
        )
    
    file_type = ALLOWED_EXTENSIONS[ext]
    
    unique_filename = f"{uuid.uuid4()}{ext}"
    file_path = os.path.join(UPLOAD_DIR, unique_filename)
    
    content = await file.read()
    file_size = len(content)
    
    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)
    
    doc = KnowledgeDocument(
        title=title,
        description=description,
        filename=file.filename,
        file_path=file_path,
        file_type=file_type,
        file_size=file_size,
        category=category,
        is_indexed=False,
        uploaded_by=current_user.id
    )
    
    db.add(doc)
    db.commit()
    db.refresh(doc)
    
    if file_type in [DocumentType.PDF.value, DocumentType.IMAGE.value]:
        background_tasks.add_task(
            index_document_smart,
            doc.id,
            file_path,
            file_type,
            title,
            category
        )
        processing_msg = "Processamento inteligente com IA em andamento."
    elif file_type in [DocumentType.PDF.value, DocumentType.DOCX.value, DocumentType.TXT.value]:
        background_tasks.add_task(
            index_document_background,
            doc.id,
            file_path,
            file_type,
            title,
            category
        )
        processing_msg = "Indexação em andamento."
    else:
        processing_msg = "Documento salvo (sem indexação de texto)."
    
    return {
        "success": True,
        "message": f"Documento enviado com sucesso. {processing_msg}",
        "document": DocumentResponse.model_validate(doc)
    }


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Remove um documento da base de conhecimento."""
    if current_user.role not in ["admin", "gestao_rv"]:
        raise HTTPException(status_code=403, detail="Acesso negado")
    
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    
    db.delete(doc)
    db.commit()
    
    return {"success": True, "message": "Documento removido com sucesso"}


@router.post("/{doc_id}/reindex")
async def reindex_document(
    doc_id: int,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """
    Reindexa um documento existente.
    PDFs e imagens utilizam automaticamente análise inteligente com GPT-4 Vision.
    """
    if current_user.role not in ["admin", "gestao_rv"]:
        raise HTTPException(status_code=403, detail="Acesso negado")
    
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
    
    if not doc:
        raise HTTPException(status_code=404, detail="Documento não encontrado")
    
    vector_store = get_vector_store()
    try:
        for i in range(doc.chunks_count or 100):
            try:
                vector_store.collection.delete(ids=[f"doc_{doc_id}_chunk_{i}"])
            except:
                pass
            try:
                vector_store.collection.delete(ids=[f"doc_{doc_id}_smart_{i}"])
            except:
                pass
    except Exception as e:
        print(f"[KNOWLEDGE] Erro ao limpar chunks antigos: {e}")
    
    doc.is_indexed = False
    doc.index_error = None
    doc.chunks_count = 0
    db.commit()
    
    if doc.file_type in [DocumentType.PDF.value, DocumentType.IMAGE.value]:
        background_tasks.add_task(
            index_document_smart,
            doc.id,
            doc.file_path,
            doc.file_type,
            doc.title,
            doc.category
        )
        msg = "Reindexação iniciada com análise inteligente"
    else:
        background_tasks.add_task(
            index_document_background,
            doc.id,
            doc.file_path,
            doc.file_type,
            doc.title,
            doc.category
        )
        msg = "Reindexação iniciada"
    
    return {"success": True, "message": msg}


@router.get("/search")
async def search_knowledge(
    query: str,
    n_results: int = 5,
    category: Optional[str] = None,
    current_user: User = Depends(get_current_user)
):
    """
    Busca semântica na base de conhecimento.
    Retorna os chunks mais relevantes para a consulta.
    """
    if not query:
        raise HTTPException(status_code=400, detail="Query é obrigatório")
    
    try:
        vector_store = get_vector_store()
        results = vector_store.search(query, n_results=n_results)
        
        if category:
            results = [r for r in results if r.get("metadata", {}).get("category") == category]
        
        return {
            "success": True,
            "query": query,
            "results": results
        }
    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "results": []
        }


@router.get("/stats")
async def get_knowledge_stats(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Retorna estatísticas da base de conhecimento."""
    if current_user.role not in ["admin", "gestao_rv"]:
        raise HTTPException(status_code=403, detail="Acesso negado")
    
    total_docs = db.query(KnowledgeDocument).count()
    indexed_docs = db.query(KnowledgeDocument).filter(KnowledgeDocument.is_indexed == True).count()
    pending_docs = db.query(KnowledgeDocument).filter(KnowledgeDocument.is_indexed == False).count()
    
    total_chunks = db.query(KnowledgeDocument).with_entities(
        db.query(KnowledgeDocument.chunks_count).scalar_subquery()
    ).scalar() or 0
    
    from sqlalchemy import func
    chunks_sum = db.query(func.sum(KnowledgeDocument.chunks_count)).scalar() or 0
    
    by_category = db.query(
        KnowledgeDocument.category,
        func.count(KnowledgeDocument.id)
    ).group_by(KnowledgeDocument.category).all()
    
    return {
        "total_documents": total_docs,
        "indexed_documents": indexed_docs,
        "pending_documents": pending_docs,
        "total_chunks": chunks_sum,
        "by_category": {cat or "Outros": count for cat, count in by_category}
    }
